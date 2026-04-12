"""
Утилиты для работы с DOCX-документами.
Модуль содержит функции для:
- настройки шрифта run-элементов;
- замены плейсхолдеров в абзацах и таблицах документа;
- добавления границ таблице;
- вставки таблицы DataFrame на место плейсхолдера;
- выделения первого абзаца жирным шрифтом.
"""

from typing import Any

import pandas as pd
from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.config import REPORT_TABLE_COLUMN_WIDTHS_INCH


def set_font(run: Run, name: str = 'Times New Roman', size: int = 11) -> None:
    """
    Устанавливает шрифт и размер для текстового фрагмента run.
    """
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)


def replace_in_paragraph(
    paragraph: Paragraph,
    placeholder: str,
    replacement_text: str,
) -> None:
    """
    Заменяет плейсхолдер в одном абзаце на переданный текст.
    Если плейсхолдер отсутствует, функция ничего не делает.
    После замены текст абзаца пересобирается через run-элементы
    с единым форматированием.
    """
    if placeholder not in paragraph.text:
        return
    full_text = paragraph.text
    parts = full_text.split(placeholder)
    paragraph.clear()
    for index, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            set_font(run)
        if index < len(parts) - 1:
            run = paragraph.add_run(replacement_text)
            set_font(run)


def replace_text_with_formatting(
    doc: DocumentType,
    placeholder: str,
    replacement_text: Any,
) -> None:
    """
    Заменяет плейсхолдер на текст во всех абзацах документа и в таблицах.
    Значение replacement_text приводится к строке перед подстановкой.
    """
    replacement_text = str(replacement_text)
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, placeholder, replacement_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, placeholder, replacement_text)


def add_table_borders(table: Table) -> None:
    """
    Добавляет границы ко всем внешним и внутренним линиям таблицы.
    """
    table_xml = table._tbl
    table_properties = table_xml.tblPr
    table_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        table_borders.append(border)
    table_properties.append(table_borders)


def add_table_at_placeholder(
    doc: DocumentType,
    df: pd.DataFrame,
    placeholder: str = '{{TABLE}}',
) -> None:
    """
    Вставляет таблицу из DataFrame на место указанного плейсхолдера.
    Плейсхолдер ищется среди абзацев документа. После нахождения
    соответствующий абзац удаляется, а на его место вставляется таблица.
    """
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            target_paragraph = paragraph
            break
    if target_paragraph is None:
        raise ValueError(f"Плейсхолдер '{placeholder}' не найден.")
    paragraph_element = target_paragraph._element
    parent = paragraph_element.getparent()
    index = parent.index(paragraph_element)
    parent.remove(paragraph_element)
    table = doc.add_table(rows=1, cols=len(df.columns))
    header_cells = table.rows[0].cells
    for column_index, column_name in enumerate(df.columns):
        cell = header_cells[column_index]
        cell.text = str(column_name)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        set_font(run, size=10)
    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for column_index, value in enumerate(row_data):
            cell = row_cells[column_index]
            cell.text = str(value)
            if column_index == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_font(cell.paragraphs[0].runs[0], size=10)
    for row in table.rows:
        for column_index, width in enumerate(REPORT_TABLE_COLUMN_WIDTHS_INCH):
            row.cells[column_index].width = Inches(width)
    table_xml = table._tbl
    table_properties = table_xml.tblPr
    table_grid = table_xml.find(qn('w:tblGrid'))
    if table_grid is not None:
        table_xml.remove(table_grid)
    table_grid = OxmlElement('w:tblGrid')
    for width in REPORT_TABLE_COLUMN_WIDTHS_INCH:
        width_dxa = int(width * 1440)
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(width_dxa))
        table_grid.append(grid_col)
    table_xml.insert(0, table_grid)
    total_width_dxa = sum(int(width * 1440) for width in REPORT_TABLE_COLUMN_WIDTHS_INCH)
    table_width = OxmlElement('w:tblW')
    table_width.set(qn('w:type'), 'dxa')
    table_width.set(qn('w:w'), str(total_width_dxa))
    table_properties.append(table_width)
    add_table_borders(table)
    table_element = table._element
    parent.insert(index, table_element)


def make_bold_first_paragraph(doc: DocumentType) -> None:
    """
    Делает все run-элементы первого абзаца документа жирными.
    """
    if not doc.paragraphs:
        return
    first_paragraph = doc.paragraphs[0]
    for run in first_paragraph.runs:
        set_font(run, name='Times New Roman', size=11)
        run.bold = True
