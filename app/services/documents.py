"""
Генерация документов акта и отчёта на основе шаблонов DOCX.
Модуль отвечает за:
- формирование акта по строке данных;
- формирование отчёта по CSV-файлу трудозатрат;
- подстановку значений в шаблоны документов;
- сохранение итоговых DOCX-файлов.
"""
from __future__ import annotations

import os
from typing import Any

import pandas as pd
from docx import Document

from app.config import ACT_TEMPLATE_FILE, OUTPUT_DIR, REPORT_TEMPLATE_FILE
from app.utils.dates import format_date, is_valid_dd_mm_yyyy, get_date_range
from app.utils.docx_utils import (
    add_table_at_placeholder,
    make_bold_first_paragraph,
    replace_text_with_formatting,
)
from app.utils.money import amount_to_words_rubles


def generate_act(row: dict[str, Any] | pd.Series, output_dir: str = OUTPUT_DIR) -> str:
    """
    Генерирует DOCX-акт по строке данных.
    В документ подставляются номер акта, даты периода и сумма вознаграждения
    в числовом и текстовом формате.
    """
    act_num = int(row['act_num'])
    month_name = row['month']
    total_amount = int(row['total_amount'])
    start_date_str = row['start_date'].strip()
    end_date_str = row['end_date'].strip()
    if not (is_valid_dd_mm_yyyy(start_date_str) and is_valid_dd_mm_yyyy(end_date_str)):
        raise ValueError(f'Неверный формат дат в строке акта №{act_num}')
    document = Document(ACT_TEMPLATE_FILE)
    total_words = amount_to_words_rubles(total_amount)
    replace_text_with_formatting(document, '{{ACT_NUM}}', act_num)
    replace_text_with_formatting(document, '{{START_DATE}}', format_date(start_date_str))
    replace_text_with_formatting(document, '{{END_DATE}}', format_date(end_date_str))
    replace_text_with_formatting(document, '{{TOTAL_AMOUNT}}', total_amount)
    replace_text_with_formatting(document, '{{TOTAL_WORDS}}', total_words)
    make_bold_first_paragraph(document)
    filename = f'Найденов Акт №{act_num} {month_name}.docx'
    path = os.path.join(output_dir, filename)
    os.makedirs(output_dir, exist_ok=True)
    document.save(path)
    print(f'✅ Акт сохранён: {filename}')
    return path


def generate_report(
    row: dict[str, Any] | pd.Series,
    output_dir: str = OUTPUT_DIR,
    debug_print: bool = False,
) -> str:
    """
    Генерирует DOCX-отчёт по строке данных и CSV-файлу трудозатрат.
    Загружает таблицу трудозатрат, рассчитывает распределение суммы
    вознаграждения по задачам, формирует итоговую таблицу и вставляет
    её в шаблон отчёта.
    """
    report_num = int(row['act_num'])
    month_name = row['month']
    total_amount = int(row['total_amount'])
    redmine_file = row['redmine_file']
    start_date_str = row['start_date'].strip()
    end_date_str = row['end_date'].strip()
    if not os.path.exists(redmine_file):
        raise FileNotFoundError(f'Файл трудозатрат не найден: {redmine_file}')
    dataframe = pd.read_csv(redmine_file, sep=';', na_values=['-', ' ', '""'])
    dataframe['Общее время'] = dataframe['Общее время'].astype(str).str.replace(',', '.', regex=True)
    dataframe['Общее время'] = pd.to_numeric(dataframe['Общее время'], errors='coerce')
    total_hours = dataframe['Общее время'].iloc[-1]
    dataframe['Расчет Вознаграждения'] = (dataframe['Общее время'] / total_hours) * total_amount
    dataframe['Расчет Вознаграждения'] = dataframe['Расчет Вознаграждения'].apply(
        lambda value: max(round(value / 50) * 50, 50)
    )
    current_total = dataframe['Расчет Вознаграждения'].iloc[:-1].sum()
    difference = total_amount - current_total
    if difference != 0:
        max_index = dataframe['Расчет Вознаграждения'].iloc[:-1].idxmax()
        dataframe.at[max_index, 'Расчет Вознаграждения'] += difference
    date_columns = dataframe.columns[1:-2]
    dataframe[['Дата начала', 'Дата окончания']] = dataframe.apply(
        lambda row_data: get_date_range(row_data, date_columns),
        axis=1,
        result_type='expand',
    )
    result_dataframe = pd.DataFrame(
        {
            '№': range(1, len(dataframe)),
            'Дата начала и окончания оказания услуги': dataframe.apply(
                lambda row_data: f"{row_data['Дата начала']} - {row_data['Дата окончания']}",
                axis=1,
            ).iloc[:-1],
            'Наименование услуги': dataframe['Задача'].iloc[:-1],
            'Расчет Вознаграждения': dataframe['Расчет Вознаграждения'].iloc[:-1],
        }
    )
    if debug_print:
        print('\n=== ТАБЛИЦА РАСЧЁТА ===')
        print(result_dataframe.to_string(index=False))
        print('========================\n')
    document = Document(REPORT_TEMPLATE_FILE)
    total_words = amount_to_words_rubles(total_amount)
    replace_text_with_formatting(document, '{{REPORT_NUM}}', report_num)
    replace_text_with_formatting(document, '{{START_DATE}}', format_date(start_date_str))
    replace_text_with_formatting(document, '{{END_DATE}}', format_date(end_date_str))
    replace_text_with_formatting(document, '{{END_DATE_SHORT}}', format_date(end_date_str, short=True))
    replace_text_with_formatting(document, '{{TOTAL_AMOUNT}}', total_amount)
    replace_text_with_formatting(document, '{{TOTAL_WORDS}}', total_words)
    add_table_at_placeholder(document, result_dataframe)
    make_bold_first_paragraph(document)
    filename = f'Найденов Отчёт №{report_num} {month_name}.docx'
    path = os.path.join(output_dir, filename)
    os.makedirs(output_dir, exist_ok=True)
    document.save(path)
    print(f'✅ Отчёт сохранён: {filename}')
    return path
